"""Aplica formato uniforme (Calibri + interlineado 1.15) al informe.docx
SIN tocar el contenido, preservando ediciones manuales."""
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DOC = Path(__file__).parent / "informe.docx"
LINE_SPACING = 1.15


def set_run_font(run, name="Calibri"):
    run.font.name = name
    # Necesario para que Word use Calibri tambien en East Asian text
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rfonts.set(qn(attr), name)


def set_paragraph_spacing(p):
    pf = p.paragraph_format
    pf.line_spacing = LINE_SPACING


def fix_paragraph(p):
    set_paragraph_spacing(p)
    for run in p.runs:
        set_run_font(run, "Calibri")


def walk_table(table):
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                fix_paragraph(p)
            for nested in cell.tables:
                walk_table(nested)


def main():
    doc = Document(DOC)

    # Default style (afecta cualquier parrafo nuevo que use Normal)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.paragraph_format.line_spacing = LINE_SPACING

    # Body paragraphs
    for p in doc.paragraphs:
        fix_paragraph(p)

    # Tables
    for table in doc.tables:
        walk_table(table)

    # Headers / footers (si los hubiera)
    for section in doc.sections:
        for p in section.header.paragraphs:
            fix_paragraph(p)
        for p in section.footer.paragraphs:
            fix_paragraph(p)

    doc.save(DOC)
    print(f"OK -> {DOC}  (Calibri, interlineado {LINE_SPACING})")


if __name__ == "__main__":
    main()
