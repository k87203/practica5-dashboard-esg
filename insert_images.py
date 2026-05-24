"""Inserta capturas de pantalla en el informe.docx en las secciones que
correspondan, con caption numerado. Preserva el contenido existente."""
from pathlib import Path

from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).parent
DOC = ROOT / "informe.docx"
LINE_SPACING = 1.15
WIDTH_CM = 15  # ancho de imagen


def set_calibri(run):
    run.font.name = "Calibri"
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rfonts.set(qn(attr), "Calibri")


def find_anchor(doc, text):
    text = text.strip()
    for p in doc.paragraphs:
        if text in p.text:
            return p
    raise RuntimeError(f"Anchor no encontrado: {text!r}")


def insert_after(anchor_p, doc, image_path, caption, num):
    pic_p = doc.add_paragraph()
    pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic_p.paragraph_format.line_spacing = LINE_SPACING
    pic_p.paragraph_format.space_before = Pt(6)
    pic_p.paragraph_format.space_after = Pt(2)
    run = pic_p.add_run()
    set_calibri(run)
    run.add_picture(str(image_path), width=Cm(WIDTH_CM))

    cap_p = doc.add_paragraph()
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_p.paragraph_format.line_spacing = LINE_SPACING
    cap_p.paragraph_format.space_after = Pt(10)
    cap_run = cap_p.add_run(f"Figura {num}. {caption}")
    cap_run.italic = True
    cap_run.font.size = Pt(10)
    set_calibri(cap_run)

    # Mover ambos justo despues del ancla, manteniendo orden [pic][cap]
    anchor_p._element.addnext(cap_p._element)
    anchor_p._element.addnext(pic_p._element)
    return cap_p  # nuevo ancla para encadenar


PLAN = [
    # (anchor_text, image_file, caption)
    ("Poblacion total.", "imagen1.PNG",
     "Vista general del dashboard: filtros, KPIs y panel principal."),
    ("Poblacion total.", "imagen6.PNG",
     "Tabla de datos filtrados exportable a CSV."),
    ("evitando redundancia.", "imagen2.PNG",
     "Scatter animado PIB vs CO2 (izquierda) y choropleth mundial (derecha)."),
    ("evitando redundancia.", "imagen3.PNG",
     "Tendencia temporal de emisiones de CO2 por continente."),
    ("costo y la latencia de la llamada.", "imagen5.PNG",
     "Datos agregados por continente enviados como contexto al modelo de lenguaje."),
    ("resumen estadistico equivalente.", "imagen4.PNG",
     "Panel de narrativa con IA: prompt editable y resumen ejecutivo generado."),
]


def main():
    doc = Document(DOC)

    # Estilo Normal con Calibri + 1.15
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.paragraph_format.line_spacing = LINE_SPACING

    # Agrupar por ancla para mantener orden de aparicion
    current_anchors = {}
    for num, (anchor_text, img, caption) in enumerate(PLAN, start=1):
        if anchor_text not in current_anchors:
            current_anchors[anchor_text] = find_anchor(doc, anchor_text)
        anchor = current_anchors[anchor_text]
        new_anchor = insert_after(anchor, doc, ROOT / img, caption, num)
        current_anchors[anchor_text] = new_anchor
        print(f"  Figura {num}: {img} -> despues de '{anchor_text[:40]}...'")

    doc.save(DOC)
    print(f"OK -> {DOC}")


if __name__ == "__main__":
    main()
