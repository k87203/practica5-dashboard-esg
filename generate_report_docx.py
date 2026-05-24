"""Genera informe.docx con el informe tecnico de la Practica 5."""
from datetime import date
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = Path(__file__).parent / "informe.docx"
REPO_URL = "https://github.com/k87203/practica5-dashboard-esg"
VIDEO_URL = "(pendiente — pegar URL de YouTube aqui)"

GREEN_DARK = RGBColor(0x0B, 0x5B, 0x43)
GREEN = RGBColor(0x11, 0x7A, 0x65)
GREY = RGBColor(0x56, 0x65, 0x73)
HEADER_BG = "117A65"
ALT_BG = "F8F9F9"


def shade(cell, color_hex: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = GREEN_DARK


def h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = GREEN


def body(doc, runs):
    """runs: list of (text, {bold, italic, code})."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for text, fmt in runs:
        r = p.add_run(text)
        r.font.size = Pt(11)
        r.bold = fmt.get("bold", False)
        r.italic = fmt.get("italic", False)
        if fmt.get("code"):
            r.font.name = "Consolas"
    return p


def para(doc, text):
    return body(doc, [(text, {})])


def bullet(doc, runs_list):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for text, fmt in runs_list:
        r = p.add_run(text)
        r.font.size = Pt(11)
        r.bold = fmt.get("bold", False)
        r.italic = fmt.get("italic", False)
        if fmt.get("code"):
            r.font.name = "Consolas"


def code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(10.5)
    # Shade the paragraph
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F4F6F7")
    p_pr.append(shd)


def styled_table(doc, header, rows, col_widths_cm):
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.autofit = False
    # Header
    hdr_cells = table.rows[0].cells
    for i, txt in enumerate(header):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        r = p.add_run(txt)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(10.5)
        shade(hdr_cells[i], HEADER_BG)
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.TOP
    # Body rows
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = ""
            p = cells[ci].paragraphs[0]
            r = p.add_run(val)
            r.font.size = Pt(10)
            cells[ci].vertical_alignment = WD_ALIGN_VERTICAL.TOP
            if ri % 2 == 1:
                shade(cells[ci], ALT_BG)
    # Column widths
    for row in table.rows:
        for i, w in enumerate(col_widths_cm):
            row.cells[i].width = Cm(w)
    return table


def build():
    doc = Document()
    # Default margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)
    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # --- PORTADA ---
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("ESG Data Explorer")
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = GREEN_DARK

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Tablero de sostenibilidad global")
    r.font.size = Pt(14)
    r.font.color.rgb = GREY

    doc.add_paragraph()
    h2(doc, "Informe Tecnico - Practica 5")
    para(doc, "Unidad 5: Visualizacion de datos")
    doc.add_paragraph()

    info = [
        ("Autor:", "Yair Garcia"),
        ("Institucion:", "Instituto Tecnologico de Tlahuac III"),
        ("Carrera:", "Ciencia de Datos"),
        ("Fecha:", date.today().strftime("%d/%m/%Y")),
        ("Repositorio:", REPO_URL),
        ("Video tutorial:", VIDEO_URL),
    ]
    t = doc.add_table(rows=len(info), cols=2)
    t.autofit = False
    for i, (k, v) in enumerate(info):
        c0, c1 = t.rows[i].cells
        c0.width = Cm(3.5); c1.width = Cm(11)
        r0 = c0.paragraphs[0].add_run(k)
        r0.bold = True
        r0.font.color.rgb = GREY
        c1.paragraphs[0].add_run(v)

    doc.add_page_break()

    # --- 1. ESCENARIO ---
    h1(doc, "1. Escenario y objetivo")
    body(doc, [
        ("El presente proyecto desarrolla un ", {}),
        ("dashboard interactivo profesional", {"bold": True}),
        (" que permite a un tomador de decisiones explorar indicadores ESG "
         "(Environmental, Social and Governance) a nivel global. "
         "Particularmente, se analiza la relacion entre el PIB y las emisiones "
         "de CO2 de los paises del mundo entre los anios 2000 y 2022, con el "
         "fin de comprender si el crecimiento economico esta acoplado o no a "
         "la huella de carbono.", {}),
    ])
    body(doc, [
        ("La fuente de datos es la ", {}),
        ("World Bank Open Data API", {"bold": True}),
        (", descargada en tiempo real mediante la libreria ", {}),
        ("requests", {"code": True}),
        (" y cacheada localmente en un CSV para garantizar reproducibilidad. "
         "Los indicadores utilizados son:", {}),
    ])
    bullet(doc, [("NY.GDP.MKTP.CD", {"code": True}),
                 (" - PIB en USD corrientes.", {})])
    bullet(doc, [("EN.GHG.CO2.MT.CE.AR5", {"code": True}),
                 (" - Emisiones de CO2 en megatoneladas.", {})])
    bullet(doc, [("SP.POP.TOTL", {"code": True}),
                 (" - Poblacion total.", {})])

    # --- 2. ARQUITECTURA ---
    h1(doc, "2. Arquitectura de la informacion")
    para(doc,
         "Cada componente visual fue elegido con un proposito narrativo "
         "especifico. La eleccion no es estetica sino comunicacional: "
         "diferentes preguntas requieren diferentes gramaticas visuales.")
    styled_table(
        doc,
        ["Visual", "Pregunta que responde", "Justificacion"],
        [
            ["KPIs (tarjetas)",
             "Cuanto y cuanto cambio?",
             "Anclaje numerico inmediato antes de explorar. Reduce carga cognitiva."],
            ["Scatter animado PIB vs CO2",
             "Como evoluciona la relacion en el tiempo?",
             "Permite observar la curva de Kuznets ambiental. La animacion comunica trayectoria, no solo estado."],
            ["Choropleth mundial",
             "Donde se concentra el problema?",
             "Mapa coropletico: intensidad geografica inmediata. Un grafico de barras con 180 paises seria ilegible."],
            ["Lineas por continente",
             "La tendencia agregada sube o baja?",
             "Series de tiempo agrupadas: tendencia macro sin perderse en el detalle pais."],
        ],
        col_widths_cm=[3.5, 5, 7.5],
    )
    doc.add_paragraph()
    body(doc, [
        ("Por que un mapa y no un grafico de barras? ", {"bold": True}),
        ("Para datos georreferenciados de ~180 entidades, el mapa aprovecha el "
         "conocimiento espacial previo del usuario: identificar Africa o Asia "
         "es instantaneo. Una barra requiere lectura secuencial de etiquetas. "
         "Ademas, el color codifica intensidad en una dimension que las barras "
         "ya usan para magnitud, evitando redundancia.", {}),
    ])

    # --- 3. METRICAS ---
    h1(doc, "3. Calculos y metricas")
    para(doc, "Las metricas reportadas en el dashboard se calculan asi:")

    h2(doc, "CO2 per capita (toneladas/habitante)")
    code_block(doc, "co2_per_capita = co2_mt * 1,000,000 / poblacion")

    h2(doc, "PIB per capita (USD/habitante)")
    code_block(doc, "gdp_per_capita = gdp_usd / poblacion")

    h2(doc, "Tasa de crecimiento interanual (%)")
    code_block(doc, "Crecimiento = ((V_actual - V_anterior) / V_anterior) * 100")
    para(doc,
         "Esta tasa se aplica al total agregado de CO2 de los paises filtrados "
         "y permite cuantificar la velocidad del cambio anio contra anio. "
         "Un valor positivo indica aceleracion de emisiones; negativo, "
         "desacople parcial del crecimiento economico.")

    h2(doc, "Agregaciones del resumen ejecutivo (IA)")
    para(doc,
         "Para el prompt enviado al modelo de lenguaje, los datos filtrados "
         "se agrupan por continente y se calcula el promedio de CO2 per "
         "capita, promedio de PIB per capita y suma total de CO2 en Mt. "
         "Esto reduce la cardinalidad antes de la inferencia y mejora el "
         "costo y la latencia de la llamada.")

    # --- 4. GUIA ---
    h1(doc, "4. Guia de usuario")
    para(doc,
         "El dashboard fue disenado para que un usuario sin formacion "
         "tecnica obtenga insights en menos de un minuto. Tres pasos:")
    bullet(doc, [
        ("Paso 1 - Filtrar. ", {"bold": True}),
        ("En la barra lateral izquierda, ajusta el rango de anios con el "
         "slider, selecciona uno o varios continentes y opcionalmente acota "
         "a paises especificos. Todos los componentes se actualizan en "
         "cascada al instante.", {}),
    ])
    bullet(doc, [
        ("Paso 2 - Explorar. ", {"bold": True}),
        ("Lee primero los cuatro KPIs superiores para tener una vision "
         "agregada. Despues, usa el scatter animado (boton play) para ver "
         "la evolucion historica y el mapa coropletico para ubicar "
         "geograficamente las emisiones del anio mas reciente.", {}),
    ])
    bullet(doc, [
        ("Paso 3 - Generar narrativa con IA. ", {"bold": True}),
        ("En el panel inferior, abre Configurar IA, pega tu API key de "
         "Anthropic y haz clic en Generar resumen. La IA producira un "
         "resumen ejecutivo de 4-6 lineas basado en los datos efectivamente "
         "filtrados. Si no cuentas con API key, el modo demo local genera un "
         "resumen estadistico equivalente.", {}),
    ])

    # --- 5. ETICA ---
    doc.add_page_break()
    h1(doc, "5. Reflexion etica")
    para(doc,
         "La visualizacion de datos es persuasiva por naturaleza y puede "
         "facilmente enganar si no se ejerce diligencia. En este proyecto "
         "se tomaron las siguientes decisiones para mitigar sesgos:")
    bullet(doc, [
        ("Escala logaritmica en PIB per capita: ", {"bold": True}),
        ("indispensable por la asimetria del dato (EEUU vs Burundi difieren "
         "por tres ordenes de magnitud). Una escala lineal aplastaria "
         "visualmente a los paises pobres. Se etiqueta el eje como "
         "'(USD, log)' para que el lector interprete correctamente.", {}),
    ])
    bullet(doc, [
        ("Per capita vs total: ", {"bold": True}),
        ("mostrar CO2 per capita evita culpar a paises poblados (China, India) "
         "por emisiones absolutas que, divididas entre su poblacion, son "
         "moderadas. Se complementa con la metrica total agregada para "
         "contextualizar el impacto global.", {}),
    ])
    bullet(doc, [
        ("Paleta secuencial (rojos) en el mapa: ", {"bold": True}),
        ("adecuada para una variable continua positiva sin punto neutro "
         "natural. Una paleta divergente (rojo-azul) sugeriria un umbral "
         "'bueno/malo' que no existe objetivamente.", {}),
    ])
    bullet(doc, [
        ("Tratamiento de datos faltantes: ", {"bold": True}),
        ("los registros con NaN en PIB o CO2 se descartan, pero se reporta "
         "el conteo de paises efectivos para que el usuario conozca la "
         "cobertura. No se imputan valores para evitar precision falsa.", {}),
    ])
    bullet(doc, [
        ("Transparencia del prompt de IA: ", {"bold": True}),
        ("el dashboard muestra explicitamente el prompt enviado al modelo. "
         "La narrativa generada queda asi auditable: el usuario puede "
         "verificar que la conclusion se deriva de los datos y no de "
         "sesgos del modelo.", {}),
    ])
    bullet(doc, [
        ("API key del usuario: ", {"bold": True}),
        ("la clave se introduce en un campo password y vive unicamente en "
         "la sesion de Streamlit; no se persiste a disco ni se transmite a "
         "terceros mas alla del proveedor del modelo.", {}),
    ])

    # --- 6. CONCLUSIONES ---
    h1(doc, "6. Conclusiones")
    body(doc, [
        ("El proyecto demuestra que es viable construir, con menos de 200 "
         "lineas de Python, una ", {}),
        ("data app", {"italic": True}),
        (" con la calidad de comunicacion esperada en un entorno profesional. "
         "Streamlit reduce la barrera entre el cientifico de datos y el "
         "usuario final, mientras que Plotly aporta la interactividad "
         "necesaria para una exploracion real (zoom, hover, animacion).", {}),
    ])
    body(doc, [
        ("La integracion de un modelo de lenguaje como capa narrativa "
         "convierte el dashboard de una herramienta de consulta en un "
         "asistente analitico: el usuario no solo ve los datos, recibe una "
         "sintesis interpretable adaptada a su seleccion. Este patron - ", {}),
        ("dashboard + LLM", {"italic": True}),
        (" - es el estandar emergente en BI moderno.", {}),
    ])

    # --- ANEXO ---
    h1(doc, "Anexo. Stack tecnologico")
    styled_table(
        doc,
        ["Componente", "Tecnologia", "Version"],
        [
            ["Lenguaje", "Python", "3.12"],
            ["Framework UI", "Streamlit", ">=1.32"],
            ["Graficos", "Plotly Express", ">=5.20"],
            ["Datos", "pandas / requests", ">=2.2 / >=2.31"],
            ["IA", "Anthropic Claude (Haiku 4.5)", ">=0.39"],
            ["Fuente", "World Bank Open Data API", "v2"],
        ],
        col_widths_cm=[4, 7, 5],
    )

    doc.save(OUT)
    print(f"OK -> {OUT}")


if __name__ == "__main__":
    build()
